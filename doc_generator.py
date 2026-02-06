"""
Word document generator for research reports.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocGenerator:
    """Generate formatted Word documents from research content."""

    def __init__(self):
        pass

    def _sanitize_filename(self, name: str) -> str:
        """
        Sanitize guest name for use in filename.

        Args:
            name: Guest name

        Returns:
            Sanitized filename
        """
        # Remove special characters and replace spaces with underscores
        sanitized = re.sub(r"[^\w\s-]", "", name)
        sanitized = re.sub(r"[-\s]+", "_", sanitized)
        return sanitized.strip("_")

    def generate_document(
        self, guest_name: str, company: str, title: str, research_content: str
    ) -> tuple[bool, str]:
        """
        Generate a formatted Word document from research content.

        Args:
            guest_name: Name of the guest
            company: Guest's company
            title: Guest's job title
            research_content: Research content from Claude

        Returns:
            tuple: (success: bool, file_path: str)
        """
        try:
            # Create document
            doc = Document()

            # Set up styles
            self._setup_styles(doc)

            # Add header
            self._add_header(doc, guest_name, company, title)

            # Add research content
            self._add_content(doc, research_content)

            # Generate filename
            sanitized_name = self._sanitize_filename(guest_name)
            filename = f"Guest_Research_{sanitized_name}.docx"
            file_path = Path(filename).absolute()

            # Save document
            doc.save(file_path)
            print(f"Document saved: {file_path}")

            return True, str(file_path)

        except Exception as e:
            print(f"Error generating document: {e}")
            return False, ""

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        # Configure normal style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_header(self, doc: Document, guest_name: str, company: str, title: str) -> None:
        """Add formatted header to document."""
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{guest_name} | {title} | {company}")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        doc.add_paragraph()

        # Add divider
        divider = doc.add_paragraph()
        divider_run = divider.add_run("_" * 80)
        divider_run.font.color.rgb = RGBColor(192, 192, 192)

        doc.add_paragraph()

    def _add_content(self, doc: Document, content: str) -> None:
        """
        Parse and add research content to document with formatting.

        Args:
            doc: Document object
            content: Research content from Claude
        """
        lines = content.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                # Add blank paragraph for spacing
                doc.add_paragraph()
                continue

            # Detect section headers (all caps or ending with colon)
            if self._is_section_header(line):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 51, 102)
                doc.add_paragraph()  # Add spacing after header
                continue

            # Detect bullet points
            if line.startswith("•") or line.startswith("-") or line.startswith("*"):
                # Remove bullet character
                text = line[1:].strip()
                para = doc.add_paragraph(text, style="List Bullet")
                continue

            # Regular paragraph
            doc.add_paragraph(line)

    def _is_section_header(self, line: str) -> bool:
        """
        Determine if a line is a section header.

        Args:
            line: Text line

        Returns:
            True if line is a header
        """
        # Check if line is short and ends with colon
        if line.endswith(":") and len(line) < 100:
            return True

        # Check if line is mostly uppercase (at least 70% uppercase)
        if len(line) > 0:
            uppercase_ratio = sum(1 for c in line if c.isupper()) / len(
                line.replace(" ", "")
            )
            if uppercase_ratio > 0.7 and len(line) < 100:
                return True

        return False
